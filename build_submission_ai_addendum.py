from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

root=Path(__file__).parent
src=root/'deliverables'/'물살림AI_대상권_절수전환_제출원고.docx'
out=root/'deliverables'/'물살림AI_대상권_최종제출원고.docx'
d=Document(src)
p=d.add_paragraph(); p.add_run().add_break()
h=d.add_paragraph(style='Heading 1'); r=h.add_run('AI 전환 근거 보강'); r.font.name='맑은 고딕'; r._element.rPr.rFonts.set(qn('w:eastAsia'),'맑은 고딕')
items=[
 '설명 가능한 전환 신호와 별도로, 탑정저수지 KRC 365일 관측에서 358개 변화 패턴을 구성해 Isolation Forest 비지도 이상 탐지 모델을 실행했다.',
 '모델 입력은 저수율, 최근 7일 기울기, 일별 절대 변화 평균이다. 최신 관측(2026-07-24)의 이상 점수는 0.505이며 과거 패턴 대비 76.0백분위다.',
 '이 값은 가뭄 예측이나 공식 경보가 아니다. 최근 변화가 해당 저수지의 과거와 얼마나 다른지 보여 주는 보조 근거로만 사용하며, 저수율·7일 추세·관측 누락·현장 정보와 함께 판단한다.',
 '모델 재현 스크립트와 산출물은 server/transition_anomaly.py, server/run_transition_anomaly.py, transition-anomaly-evaluation.json에 보관한다. 성능 미달 7일 예측 후보를 배포하지 않는 기존 거버넌스 원칙도 그대로 유지한다.'
]
for x in items:
 p=d.add_paragraph(style='List Bullet'); r=p.add_run(x); r.font.name='맑은 고딕'; r._element.rPr.rFonts.set(qn('w:eastAsia'),'맑은 고딕'); r.font.size=Pt(10.5)
d.save(out); print(out)
