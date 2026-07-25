from pathlib import Path
from docx import Document

root=Path(__file__).parent
source=root/'deliverables'/'물살림AI_실사용중심_최종제출원고_v2.docx'
output=root/'deliverables'/'물살림AI_실사용중심_최종제출원고_릴리스.docx'
replacements={
    'mulsallim-public-v2.html':'mulsallim-public-release.html',
    'mulsallim-dashboard-practice.html':'mulsallim-dashboard-release.html',
}
d=Document(source)
def swap(paragraph):
    for run in paragraph.runs:
        for before,after in replacements.items():
            if before in run.text: run.text=run.text.replace(before,after)
for p in d.paragraphs: swap(p)
for table in d.tables:
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs: swap(p)
d.save(output)
print(output)
