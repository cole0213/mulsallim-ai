from pathlib import Path
from docx import Document

root=Path(__file__).parent
source=root/'deliverables'/'물살림AI_실사용중심_최종제출원고_릴리스.docx'
output=root/'deliverables'/'물살림AI_실사용중심_최종제출원고_위치기반.docx'
replacements={
    'http://localhost:4182/mulsallim-public-release.html':'http://localhost:4183/mulsallim-public-county.html',
    'http://localhost:4182/mulsallim-dashboard-release.html':'http://localhost:4183/mulsallim-dashboard-county.html',
    '충청남도 → 논산시 → 탑정 선택':'충청남도 → 논산시 → KRC 관측 저수지 선택',
}
d=Document(source)
def swap(p):
    for run in p.runs:
        for before,after in replacements.items():
            run.text=run.text.replace(before,after)
for p in d.paragraphs:swap(p)
for t in d.tables:
    for row in t.rows:
        for c in row.cells:
            for p in c.paragraphs:swap(p)
d.save(output)
print(output)
