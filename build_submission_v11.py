from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

ROOT=Path(__file__).parent
SRC=ROOT/'deliverables'/'물살림AI_대상권_실증단계_최종제출원고_v10.docx'
OUT=ROOT/'deliverables'/'물살림AI_대상권_실증·안전성_최종제출원고_v11.docx'

def font(run,size=10.2,bold=None):
    run.font.name='맑은 고딕';run._element.rPr.rFonts.set(qn('w:eastAsia'),'맑은 고딕');run.font.size=Pt(size)
    if bold is not None:run.bold=bold
def heading(doc,text):
    p=doc.add_paragraph();p.paragraph_format.space_before=Pt(16);p.paragraph_format.space_after=Pt(6);font(p.add_run(text),15,True)
def bullet(doc,text):
    p=doc.add_paragraph(style='List Bullet');p.paragraph_format.space_after=Pt(3);font(p.add_run(text),10.2)

d=Document(SRC)
heading(d,'13. 개인정보 최소화와 서버 안전성')
p=d.add_paragraph();font(p.add_run('물살림 AI는 개인정보를 받지 않는다는 문구에만 의존하지 않는다. 협의 카드와 익명 파일럿 피드백의 저장 요청에 전화번호·이메일 형식이 포함되면 서버가 저장 전 HTTP 400으로 거부한다.'),10.2)
bullet(d,'협의 카드는 저수지명·일반 지역명·관측일·판단·확인 항목만 30일간 보관한다.')
bullet(d,'파일럿 피드백은 역할 분류와 1~5점 사용성 문항, 선택 의견만 90일간 보관한다.')
bullet(d,'공개 협의 카드는 읽기 전용이고, 상태 변경은 생성자에게만 발급된 비공개 키가 있어야 한다.')
bullet(d,'2026-07-24 회귀 검증에서 전화번호·이메일 포함 카드와 피드백이 서버에서 저장 거부됨을 확인했다.')
heading(d,'14. 최종 시연·제출 기준')
p=d.add_paragraph();font(p.add_run('기본 진입: http://localhost:4184/  |  대시보드: mulsallim-dashboard-county-v10.html  |  다지역 모델카드: mulsallim-ai-model-card-v3.html  |  파일럿 피드백: mulsallim-pilot.html  |  서버 안전성: server-production-v11.js'),10.2)
d.save(OUT);print(OUT)
