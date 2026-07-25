from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

ROOT=Path(__file__).parent
SRC=ROOT/'deliverables'/'물살림AI_대상권_최종제출원고_v7.docx'
OUT=ROOT/'deliverables'/'물살림AI_대상권_실증단계_최종제출원고_v10.docx'

def font(run,size=10.2,bold=None):
    run.font.name='맑은 고딕';run._element.rPr.rFonts.set(qn('w:eastAsia'),'맑은 고딕');run.font.size=Pt(size)
    if bold is not None:run.bold=bold
def replace_runs(paragraph):
    for run in paragraph.runs:
        run.text=run.text.replace('v7.html','v10.html').replace('v7','v10')
def add_heading(doc,text):
    p=doc.add_paragraph();p.paragraph_format.space_before=Pt(16);p.paragraph_format.space_after=Pt(6);font(p.add_run(text),15,True)
def add_bullet(doc,text):
    p=doc.add_paragraph(style='List Bullet');p.paragraph_format.space_after=Pt(3);font(p.add_run(text),10.2)

d=Document(SRC)
for p in d.paragraphs:replace_runs(p)
for table in d.tables:
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:replace_runs(p)

add_heading(d,'10. 다지역 시간 분리 AI 검증과 실증 체계')
intro=d.add_paragraph();font(intro.add_run('물살림 AI는 한 저수지의 결과로 AI 성능을 일반화하지 않았다. 9개 권역 18개 저수지의 최근 1년 KRC 일별 관측 6,588건을 수집하고, 저수지별 과거 70%로 학습한 뒤 미래 30%(1,908개 예시)로 홀드아웃 평가했다.'),10.2)
add_bullet(d,'7일 저수율 예측에서 현재값 유지 기준선 MAE는 3.525%p, 기울기 투영 후보는 4.198%p였다. 후보가 더 나빠 예측값을 배포하지 않는다.')
add_bullet(d,'7일 내 5%p 이상 하락 확인 신호는 저수율 50% 이하 기준 F1 0.040 대비 세 특징 후보 F1 0.062로 개선됐다. 그러나 재현율 0.037로 자동 알림·공식 경보·급수 결정을 위한 성능은 아니다.')
add_bullet(d,'따라서 현재 서비스는 설명 가능한 현장 확인 우선순위만 제공하며, 모델카드에서 범위·성능·보류 근거를 공개한다.')
add_heading(d,'11. 파일럿 측정 기능과 운영 원칙')
add_bullet(d,'익명 파일럿 화면은 농업인·수리계, KRC 현장 담당자, 지자체 담당자의 사용 편의·이해도·재사용 의향을 1~5점으로 기록한다.')
add_bullet(d,'설문 응답은 90일, 협의 카드는 30일 뒤 자동 정리한다. 이름·연락처·정확한 농지 주소는 저장하지 않는다.')
add_bullet(d,'응답 수가 적으면 성과로 포장하지 않고 표본 수·역할·기간·불편 의견을 함께 공개한다.')
add_heading(d,'12. 최신 시연 경로')
p=d.add_paragraph();font(p.add_run('기본 진입: http://localhost:4184/  |  대시보드: mulsallim-dashboard-county-v10.html  |  다지역 모델카드: mulsallim-ai-model-card-v3.html  |  파일럿 피드백: mulsallim-pilot.html'),10.2)
d.save(OUT);print(OUT)
