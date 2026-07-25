from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn

ROOT=Path(__file__).parent
OUT=ROOT/'deliverables'/'물살림AI_대상권_절수전환_제출원고.docx'
doc=Document(); sec=doc.sections[0]
sec.top_margin=sec.bottom_margin=Cm(1.8); sec.left_margin=sec.right_margin=Cm(2.0)
for style in ['Normal','Heading 1','Heading 2']:
    s=doc.styles[style]; s.font.name='맑은 고딕'; s._element.rPr.rFonts.set(qn('w:eastAsia'),'맑은 고딕')
doc.styles['Normal'].font.size=Pt(10.5)
def add(text='',style=None,bold=False,size=None,color=None):
    p=doc.add_paragraph(style=style); r=p.add_run(text); r.bold=bold; r.font.name='맑은 고딕'; r._element.rPr.rFonts.set(qn('w:eastAsia'),'맑은 고딕')
    if size:r.font.size=Pt(size)
    if color:r.font.color.rgb=RGBColor(*color)
    return p
add('물살림 AI',bold=True,size=26,color=(8,112,109)); add('절수 전환 의사결정 서비스',bold=True,size=16); add('제3회 KRC AI 디지털 혁신 공모전 · 서비스개발 분야 제출 원고',size=10,color=(89,112,107))
sections=[
('1. 서비스 한 줄 정의','물살림 AI는 KRC 저수지 수위 원자료를 바탕으로 누가, 언제, 어떤 근거로 절수 협의를 시작해야 하는지를 설명하고, 절수 선택의 전환 지연 효과를 지역 구성원이 함께 검토·공유하게 하는 실행형 의사결정 서비스다.'),
('2. 문제와 차별성','ADMS 농업가뭄관리시스템은 이미 저수율 현황·지도·공식 예경보·저수율 전망을 제공한다. 물살림 AI는 이를 대체하지 않는다. 대신 “정보가 있는데도 절수 협의는 왜 늦어지는가”라는 실행 공백을 해결한다. 사용자는 이번 주 협의 필요성, 관심·주의선 도달 추정, 10%·20% 절수에 따른 지연 효과, 같은 시·군의 우선 협의 대상을 한 화면에서 확인한다.'),
('3. 실제 구현 기능','① 시·도→시·군·구→저수지 선택: KRC reservoircode로 후보를 찾고 시설 코드는 화면에 노출하지 않는다.\n② KRC 원자료 분석: reservoirlevel의 관측일·수위·저수율을 30·60·90일 축·점 도구 설명·경계선과 함께 표시한다.\n③ AI 전환 신호: 저수율, 최근 7일 하락, 과거 변동성 대비 이상도, 관측 누락을 합산해 협의 필요성을 0~100으로 설명한다.\n④ 절수 선택: 최근 강건 추세가 유지된다는 탐색 가정 아래 0·10·20% 절수의 주의선 도달 시점을 비교한다.\n⑤ 실행: 동일 시·군 우선순위, 관심 저수지 저장, 공유용 TXT 브리핑을 제공한다.'),
('4. 설명 가능한 AI와 검증','전환 신호는 최근 관측점의 쌍별 기울기 중앙값(Theil–Sen 방식)으로 하락 추세를 구하고, 과거 일별 변화의 중앙값·MAD(중앙절대편차)와 비교해 이상도를 계산한다. 관측 누락이 있으면 신뢰도를 낮춘다. 이는 공식 경보·확정 예측이 아니라 현장 협의 시작을 돕는 탐지 지표다.\n\n탑정저수지 365일 데이터의 7일 예측에서 전일값 유지 기준모델은 MAE 9.817%p, MAPE 18.913%였고 Gradient Boosting 후보는 MAE 18.633%p, MAPE 41.386%였다. 후보가 기준보다 낮아 배포하지 않았다. 강수량·가뭄 단계·유효저수량·수혜면적·작물 수요를 결합한 다변량 후보가 동일 검증을 통과할 때만 예측 기능을 켠다.'),
('5. 공익성·사업성','초기에는 KRC 지사 또는 지자체 1곳, KRC 관리 저수지 5~10곳, 농업인 대표 그룹을 대상으로 8주 파일럿을 운영한다. 위험 신호부터 첫 협의까지 걸린 시간, 상위 우선순위와 담당자 판단의 일치율, 시나리오 열람 뒤 협의·점검 기록 비율, 판단 이해도 설문을 KPI로 측정한다. 이후 지역 맞춤 브리핑·알림·현장 점검 이력 기능으로 확장한다.'),
('6. 신뢰·안전 원칙','인증키는 브라우저·URL·문서에 넣지 않고 서버 환경변수에만 보관한다. API 오류·빈 응답을 그럴듯한 예시값으로 대체하지 않는다. 저수율·시나리오·전환 신호는 ADMS 공식 예경보, 현장 수요, 강수·유입량, 수리시설 상태를 대체하지 않으며 모든 화면과 브리핑에 이 한계를 표시한다.'),
('7. 시연 순서','1) 랜딩에서 “ADMS가 정보와 전망을 제공한다면, 협의는 누가 언제 시작할까?”를 제시한다.\n2) 지역·저수지를 선택해 KRC 원자료와 신뢰 상태를 보여 준다.\n3) AI 전환 신호와 근거 카드로 이번 주 판단을 설명한다.\n4) 관심·주의선 도달 추정과 10%·20% 절수 지연 효과를 비교한다.\n5) 동일 지역 우선순위와 TXT 브리핑으로 행동까지 연결한다.\n6) 성능 미달 예측 후보를 배포하지 않은 거버넌스로 책임 있는 AI를 마무리한다.'),
]
for h,body in sections:
    add(h,'Heading 1')
    for para in body.split('\n'):
        add(para)
add('8. 심사 기준 대응','Heading 1')
tbl=doc.add_table(rows=1,cols=2); tbl.style='Table Grid'
for cell,text in zip(tbl.rows[0].cells,['평가 항목','대상권 근거']):
    cell.text=text
    for r in cell.paragraphs[0].runs:r.bold=True;r.font.name='맑은 고딕';r._element.rPr.rFonts.set(qn('w:eastAsia'),'맑은 고딕')
for x,y in [('활용성 30','KRC 원자료를 탐색부터 전환 판단·시나리오·공유까지 전 흐름에 사용'),('창의성 20','조회·전망의 중복이 아닌 절수 협의 시작 시점이라는 실행 공백 해결'),('사업성 20','5~10개 저수지 8주 파일럿, KPI, KRC→지자체→농업인 확장 경로'),('AI 활용성 20','데이터 적응형 전환 탐지, 설명 가능한 근거, 성능 미달 후보 배제'),('완성도 10','실제 API·지역 선택·차트·한계·비교·저장·공유·오류 처리')]:
    row=tbl.add_row().cells; row[0].text=x; row[1].text=y
    for c in row:
        for p in c.paragraphs:
            for r in p.runs:r.font.name='맑은 고딕';r._element.rPr.rFonts.set(qn('w:eastAsia'),'맑은 고딕')
add('최종 시연 링크','Heading 1'); add('랜딩: http://localhost:4182/mulsallim-home.html'); add('대시보드: http://localhost:4182/mulsallim-action.html')
OUT.parent.mkdir(exist_ok=True); doc.save(OUT); print(OUT)
