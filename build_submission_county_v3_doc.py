from pathlib import Path
from docx import Document

root=Path(__file__).parent
source=root/'deliverables'/'물살림AI_실사용중심_최종제출원고_위치기반.docx'
output=root/'deliverables'/'물살림AI_실사용중심_최종제출원고_최종v3.docx'
d=Document(source)
for p in d.paragraphs:
    for r in p.runs:
        r.text=r.text.replace('mulsallim-public-county.html','mulsallim-public-county-v3.html').replace('mulsallim-dashboard-county.html','mulsallim-dashboard-county-v3.html')
for t in d.tables:
    for row in t.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:r.text=r.text.replace('mulsallim-public-county.html','mulsallim-public-county-v3.html').replace('mulsallim-dashboard-county.html','mulsallim-dashboard-county-v3.html')
d.save(output)
print(output)
