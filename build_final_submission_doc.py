from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn
root=Path(__file__).parent; src=root/'docs'/'08_최종_제출_명세서_및_원고.md'; out=root/'deliverables'/'물살림AI_최종_서비스개발_제출원고.docx'
d=Document(); s=d.sections[0]; s.top_margin=s.bottom_margin=Cm(2); s.left_margin=s.right_margin=Cm(2.1)
d.styles['Normal'].font.name='맑은 고딕'; d.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'),'맑은 고딕'); d.styles['Normal'].font.size=Pt(10.5)
p=d.add_paragraph(); r=p.add_run('물살림 AI\n최종 서비스개발 제출원고'); r.bold=True; r.font.size=Pt(23)
for raw in src.read_text(encoding='utf-8').splitlines():
    x=raw.strip()
    if not x or x.startswith('>') or x.startswith('# '): continue
    if x.startswith('## '): d.add_paragraph(x[3:],style='Heading 1')
    elif x.startswith('- '): d.add_paragraph(x[2:],style='List Bullet')
    elif x.startswith('1. **') or x[:2].isdigit() and '. ' in x: d.add_paragraph(x.replace('**',''))
    else: d.add_paragraph(x.replace('**','').replace('`',''))
d.save(out); print(out)
