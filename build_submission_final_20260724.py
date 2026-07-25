from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn

ROOT = Path(__file__).parent
OUT = ROOT / 'deliverables' / '물살림AI_제출용_최종_서비스명세서_20260724.docx'

def kr(run, name='맑은 고딕'):
    run.font.name = name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)

def paragraph(doc, text='', style=None, bold=False, size=None, color=None):
    p = doc.add_paragraph(style=style)
    r = p.add_run(text)
    r.bold = bold
    kr(r)
    if size: r.font.size = Pt(size)
    if color: r.font.color.rgb = RGBColor(*color)
    return p

d = Document()
s = d.sections[0]
s.top_margin = Cm(1.8); s.bottom_margin = Cm(1.8); s.left_margin = Cm(2.0); s.right_margin = Cm(2.0)
normal = d.styles['Normal']; normal.font.name = '맑은 고딕'; normal._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕'); normal.font.size = Pt(10.5)
for st in ['Heading 1', 'Heading 2']:
    d.styles[st].font.name = '맑은 고딕'; d.styles[st]._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

paragraph(d, '물살림 AI', bold=True, size=26, color=(8,112,109))
paragraph(d, '지역 농업용수 위험 판단 및 절수 협의 지원 서비스', bold=True, size=16)
paragraph(d, '서비스개발 분야 제출용 최종 서비스 명세서 | 2026. 07. 24.', size=10, color=(89,112,107))

paragraph(d, '1. 서비스 개요', 'Heading 1')
paragraph(d, '물살림 AI는 한국농어촌공사(KRC) 농촌용수 저수지 수위 공공데이터를 국민·농업인·현장 담당자가 바로 이해할 수 있는 지역 기반 의사결정 화면으로 바꾸는 서비스다. 단순 수치 조회가 아니라 저수율, 최근 7일 변화, 위험 경계선, 관측 결측 상태를 함께 제시하고 동일 지역 비교와 작물별 절수 시나리오까지 연결한다.')

paragraph(d, '2. 문제 인식과 해결', 'Heading 1')
paragraph(d, '저수지 수위 정보는 존재하지만 주민과 농업인은 시설 코드와 개별 수치만으로 지금의 위험도와 다음 행동을 판단하기 어렵다. 물살림 AI는 지역 선택부터 시작해 사용자가 시설 코드를 보지 않아도 관련 저수지를 찾도록 한다. 관측 원자료를 추세·경계선·근거 문장으로 해석하고, 현장 협의에 필요한 비교와 절수 선택지를 제공한다.')

paragraph(d, '3. 실제 배포 대시보드 기능', 'Heading 1')
for t in [
    '지역 중심 탐색: 시·도와 시·군·구를 고르면 KRC 저수지 후보를 실시간 조회한다. 시설 코드는 내부 요청에만 쓰며 화면에 노출하지 않는다.',
    '인터랙티브 추세: 7·30·90일 전환, X·Y축 수치, 50% 관심선과 35% 주의선, 관측점 도구 설명으로 저수율·수위·관측일을 확인한다.',
    '설명 가능한 위험 판단: 저수율 35% 이하 또는 최근 7일 15%p 이상 하락은 주의, 저수율 50% 이하 또는 최근 7일 7%p 이상 하락은 관심으로 표시한다. 적용된 기준과 관측 건수·누락 건수·최신 관측일을 함께 보여 준다.',
    '동일 지역 비교: 같은 시·군·구 저수지를 최신 저수율과 7일 변화로 정렬해 우선 확인 대상을 찾는다.',
    '작물별 절수 시나리오: 논·밭작물·과수, 생육 단계, 면적을 입력하면 기본·10%·20% 절수 시 일일 예상 용수 수요를 비교한다. 이는 현장 협의용 추정치임을 명시한다.'
]:
    p = d.add_paragraph(style='List Bullet'); r=p.add_run(t); kr(r)

paragraph(d, '4. AI 활용과 검증 원칙', 'Heading 1')
paragraph(d, '서버는 매일 KRC 관측치를 수집하고 기준모델과 후보모델을 같은 검증 방식으로 비교한다. 탑정저수지 365일 데이터의 7일 예측 검증에서 전일값 유지 기준모델은 MAE 9.817%p, MAPE 18.913%였고, Gradient Boosting 후보는 MAE 18.633%p, MAPE 41.386%였다. 후보 모델이 기준보다 나빠 배포 모델로 채택하지 않았다. 강수량·가뭄 단계·용수 수요 특성을 결합한 뒤 동일 기준으로 재검증할 계획이며, 현재 대시보드에도 이 배제 사유를 투명하게 표시한다.')

paragraph(d, '5. 공공데이터·보안·신뢰성', 'Heading 1')
for t in [
    '데이터 원천: 한국농어촌공사 농촌용수 저수지 수위정보 API의 reservoircode, reservoirlevel.',
    '인증키는 브라우저·URL·문서에 노출하지 않고 서버 환경변수에만 보관한다.',
    'API 오류·빈 응답은 예시 수치로 대체하지 않고 오류 원인과 재시도 필요성을 안내한다.',
    '위험 표시는 법정 가뭄 경보, 급수 지시, 농업용수 배분 결정을 대체하지 않는 의사결정 보조 정보다.'
]:
    p=d.add_paragraph(style='List Bullet'); r=p.add_run(t); kr(r)

paragraph(d, '6. 심사 기준 대응', 'Heading 1')
table=d.add_table(rows=1, cols=2); table.style='Table Grid'
for cell, txt in zip(table.rows[0].cells, ['평가 항목', '구현 근거']):
    cell.text=txt
    for r in cell.paragraphs[0].runs: r.bold=True; kr(r)
for a,b in [
    ('활용성 (30)', 'KRC 원자료를 지역 선택·추세·비교·절수 협의에 직접 활용'),
    ('창의성 (20)', '수위 조회를 지역 기반 위험 근거와 공동 절수 협의 흐름으로 전환'),
    ('사업성 (20)', '주민·농업인·현장 담당자가 동일 화면을 사용하고 지역별 확장 가능'),
    ('AI 활용성 (20)', '성능 수치 검증, 미달 후보 배제, 재검증 계획을 공개'),
    ('완성도 (10)', '랜딩과 실제 도구 분리, 축·도구 설명·오류·결측·관측일·면책 표시')]:
    cells=table.add_row().cells; cells[0].text=a; cells[1].text=b
    for c in cells:
        for p in c.paragraphs:
            for r in p.runs: kr(r)

paragraph(d, '7. 시연 링크와 발표 순서', 'Heading 1')
paragraph(d, '랜딩 페이지: http://localhost:4182/mulsallim.html')
paragraph(d, '실제 대시보드: http://localhost:4182/mulsallim-dashboard.html')
for t in ['랜딩에서 목적과 핵심 기능을 소개한다.', '지역을 고르고 저수지를 선택한다.', '추세 그래프와 위험 근거를 확인한다.', '동일 지역 비교와 절수 시나리오를 제시한다.', '모델 검증 결과와 배포 제외 근거를 설명한다.']:
    p=d.add_paragraph(style='List Number'); r=p.add_run(t); kr(r)

OUT.parent.mkdir(exist_ok=True)
d.save(OUT)
print(OUT)
