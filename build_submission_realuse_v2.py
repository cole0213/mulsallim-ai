from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT=Path(__file__).parent
OUT=ROOT/'deliverables'/'물살림AI_실사용중심_최종제출원고_v2.docx'
INK='143B3B'; TEAL='087A77'; MINT='E7F3EE'; LINE='C9DDD4'

def shade(cell, fill):
    tcPr=cell._tc.get_or_add_tcPr(); shd=OxmlElement('w:shd'); shd.set(qn('w:fill'),fill); tcPr.append(shd)
def border(cell, color=LINE):
    tcPr=cell._tc.get_or_add_tcPr(); b=OxmlElement('w:tcBorders')
    for side in ('top','left','bottom','right'):
        e=OxmlElement(f'w:{side}'); e.set(qn('w:val'),'single'); e.set(qn('w:sz'),'4'); e.set(qn('w:color'),color); b.append(e)
    tcPr.append(b)
def para(doc, text='', style=None, bold=False, color=None, size=None, before=None, after=None):
    p=doc.add_paragraph(style=style); r=p.add_run(text); r.bold=bold
    if color:r.font.color.rgb=RGBColor.from_string(color)
    if size:r.font.size=Pt(size)
    if before is not None:p.paragraph_format.space_before=Pt(before)
    if after is not None:p.paragraph_format.space_after=Pt(after)
    return p
def bullet(doc, text):
    p=doc.add_paragraph(style='List Bullet'); p.add_run(text); return p
def heading(doc, text, level=1):
    p=doc.add_paragraph(style=f'Heading {level}'); p.add_run(text); return p
def table(doc, headers, rows, widths=None):
    t=doc.add_table(rows=1, cols=len(headers)); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.style='Table Grid'
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; c.text=h; shade(c,TEAL); border(c,'FFFFFF'); c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for r in c.paragraphs[0].runs:r.font.bold=True;r.font.color.rgb=RGBColor(255,255,255)
    for row in rows:
        cells=t.add_row().cells
        for i,val in enumerate(row):
            cells[i].text=val; shade(cells[i],'F8FCFA' if len(t.rows)%2 else 'FFFFFF'); border(cells[i]); cells[i].vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cells[i].paragraphs:p.paragraph_format.space_after=Pt(3);p.paragraph_format.space_before=Pt(3)
    if widths:
        for row in t.rows:
            for i,w in enumerate(widths):row.cells[i].width=Cm(w)
    doc.add_paragraph().paragraph_format.space_after=Pt(2)
    return t

d=Document(); sec=d.sections[0]; sec.top_margin=Cm(1.8);sec.bottom_margin=Cm(1.8);sec.left_margin=Cm(2.0);sec.right_margin=Cm(2.0)
normal=d.styles['Normal'];normal.font.name='맑은 고딕';normal._element.rPr.rFonts.set(qn('w:eastAsia'),'맑은 고딕');normal.font.size=Pt(10.3);normal.font.color.rgb=RGBColor.from_string(INK);normal.paragraph_format.space_after=Pt(6)
for level,size in [(1,18),(2,13)]:
    s=d.styles[f'Heading {level}'];s.font.name='맑은 고딕';s._element.rPr.rFonts.set(qn('w:eastAsia'),'맑은 고딕');s.font.size=Pt(size);s.font.bold=True;s.font.color.rgb=RGBColor.from_string(TEAL);s.paragraph_format.space_before=Pt(18);s.paragraph_format.space_after=Pt(8)

p=para(d,'제3회 KRC AI 디지털 혁신 공모전 | 서비스개발 분야',bold=True,color=TEAL,size=11,after=15)
p=para(d,'물살림 AI',bold=True,color=INK,size=29,after=2);p.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.LEFT
para(d,'관측값을 현장 행동으로 연결하는 농업용수 재확인 서비스',bold=True,color=TEAL,size=16,after=18)
lead=para(d,'제안 요약  KRC 농촌용수 저수지 수위정보를 기반으로, 농업인과 현장 담당자가 “우리 지역 물 상황이 평소와 다른가?”, “이번 주에 무엇을 확인해야 하는가?”를 코드 없이 확인하고, 근거와 현장 메모를 공유하도록 돕는 웹 서비스입니다.',size=11,after=8)
lead.paragraph_format.line_spacing=1.45
table(d,['평가 관점','제출 핵심'],[
 ['활용성','KRC 공개 저수지 코드·수위정보를 실제 지역 선택과 관측 근거 화면에 연결'],
 ['창의성','공식 경보를 흉내 내지 않고 “절수 협의 시작 시점”과 현장 대화에 집중'],
 ['사업성','관심 저수지·재확인 일정·공유 요약으로 1회 조회가 아닌 반복 사용을 설계'],
 ['AI 활용성','검증 통과 모델만 반영하고, 미달 예측 모델은 배포에서 제외하는 거버넌스 적용'],
 ['완성도','지역→저수지→근거→이번 주 행동→공유·재확인의 단일 흐름 구현'],[3,12])

heading(d,'1. 해결하려는 실제 문제')
para(d,'저수율은 공개되어 있어도 사용자는 자신의 지역 저수지를 찾고, 숫자의 변화를 해석하고, 누구와 무엇을 논의할지까지 직접 연결해야 합니다. 기존 정보 화면은 현황 조회에는 강하지만, 농업인과 현장 담당자의 반복 과업인 “변화를 놓치지 않고 현장 확인을 시작하는 일”을 한 흐름으로 만들지 못합니다.')
para(d,'물살림 AI는 예측 수치 경쟁이 아니라, 관측값을 이해 가능한 행동으로 번역하는 의사결정 보조 레이어를 제안합니다.',bold=True,color=TEAL)

heading(d,'2. 사용자와 핵심 과업')
table(d,['사용자','실제 질문','서비스가 답하는 방식'],[
 ['농업인','우리 지역 물 상황이 평소와 다른가?','최신 관측일·저수율·7일 변화와 관측 공백을 함께 확인'],
 ['농업인/이용자','이번 주에 무엇을 해야 하나?','정기 관측·절수 준비·현장 점검의 보조 판단과 체크리스트 제공'],
 ['현장 담당자','어떤 근거로 대화를 시작할까?','판단 근거 카드, 현장 메모, 근거 포함 공유 요약 제공'],
 ['관리자','다시 확인할 시점을 어떻게 놓치지 않을까?','관심 저수지 저장과 캘린더 재확인 일정(.ics) 제공'],[3,6,6])

heading(d,'3. 서비스 흐름과 실제 구현')
table(d,['순서','사용자 경험','구현 기능'],[
 ['01 지역 선택','시·도와 시·군·구만 고른다.','사용자에게 저수지 코드를 노출하지 않고 KRC 코드 조회 API로 대상 목록 생성'],
 ['02 저수지 선택','이름과 위치를 보고 대상 저수지를 고른다.','KRC 수위 API에서 최근 30/60/90일 관측값 호출'],
 ['03 근거 확인','현재 저수율과 변화가 왜 중요한지 본다.','저수율, 7일 변화, 하락 패턴, 관측 공백·신뢰도, 그래프 툴팁 제공'],
 ['04 이번 주 행동','화면을 닫기 전에 확인할 일을 정한다.','위험 단계별 현장 점검 체크리스트, ADMS 공식 정보 링크, 기기 내 메모 제공'],
 ['05 공유·재확인','담당자와 같은 근거를 보고 다음 확인을 예약한다.','TXT 브리핑, 관심 저수지 저장, 캘린더 일정 파일 제공'],[1.3,4.6,9.1])

heading(d,'4. AI·데이터 활용성과 안전한 한계')
para(d,'데이터는 한국농어촌공사 농촌용수 저수지 수위정보의 저수지명·위치·관측일·수위·저수율을 사용합니다. 화면에는 데이터 출처, 관측기간, 누락일 수, 신뢰도 점수를 노출해 사용자가 수치의 상태를 함께 판단할 수 있게 했습니다.')
para(d,'AI는 법정 가뭄경보를 발령하거나 급수·절수 결정을 자동화하지 않습니다. 최근 7일의 강건 추세와 과거 일별 변화 분포를 비교해 하락 패턴의 이상성을 보조 근거로 제시합니다. Isolation Forest 기반 전환 이상성 평가의 최신 결과는 2026-07-24 관측 기준 이상성 점수 0.505, 과거 분포 백분위 76.0입니다. 이 수치는 가뭄 예보가 아니라 “과거 패턴과 얼마나 다른가”를 설명하는 근거입니다.')
para(d,'또한 Gradient Boosting 예측 후보는 기준선보다 성능이 낮았습니다(MAE 18.633%p, 기준선 9.817%p). 따라서 배포 판단에서 제외했습니다. 모델이 있다는 주장보다, 검증에서 이기지 못한 모델을 제외하는 원칙이 공공 서비스의 신뢰를 높인다고 판단했습니다.',bold=True)
table(d,['서비스가 하는 일','서비스가 하지 않는 일'],[
 ['관측 변화를 정리하고 현장 확인·공유를 돕는다.','공식 가뭄경보, 급수 배분, 절수 지시를 대체하지 않는다.'],
 ['근거와 데이터 품질을 함께 보여 준다.','불확실한 예측값을 확정값처럼 제시하지 않는다.'],
 ['ADMS 공식 정보를 함께 확인하도록 연결한다.','권한 없는 자동 의사결정을 실행하지 않는다.'],[7.5,7.5])

heading(d,'5. 차별성 및 확장성')
para(d,'ADMS와 같은 공식 정보 체계는 가뭄 현황과 안내의 기준입니다. 물살림 AI는 이를 대체하려 하지 않습니다. 차별점은 개별 사용자의 지역·저수지 단위에서 관측값을 “오늘 무엇을 확인할지”로 번역하고, 현장 메모·공유·재확인이라는 반복 사용 흐름을 만드는 데 있습니다.')
bullet(d,'1단계: KRC 공개 수위정보 기반 지역 선택·근거·체크리스트·공유 기능 운영')
bullet(d,'2단계: 사용자 동의 기반 현장 확인 기록을 익명·집계하여 지역 담당자 협의 자료로 확장')
bullet(d,'3단계: KRC와 협의한 추가 데이터(유효저수량, 수혜면적, 시설 상태 등)가 확보될 때만 검증된 모델 범위를 확장')

heading(d,'6. 성과 측정과 현장 검증 계획')
table(d,['지표','측정 방법','목표/판단 기준'],[
 ['반복 사용','관심 저수지 저장 후 재방문 비율','현장 파일럿에서 주간 재확인 행동이 일어나는지 검증'],
 ['이해 가능성','근거 카드 후 “다음 행동” 선택 성공률','농업인·담당자가 같은 판단 근거를 설명 가능한지 확인'],
 ['공유 실용성','브리핑 복사·다운로드·메모 사용 건수','수치 조회가 현장 대화로 이어지는지 확인'],
 ['안전성','공식 안내 링크 확인 및 오판 신고','공식 경보 대체로 오인되는 표현·흐름을 지속 개선'],[3,6,6])
para(d,'현재 파일럿 성과를 이미 보유했다고 주장하지 않습니다. 제출 후 KRC·지자체 담당자 및 농업인 5~10명을 대상으로 사용성 검증을 수행하고, 반복 사용·이해·공유·안전성 지표를 바탕으로 개선합니다.',bold=True,color=TEAL)

heading(d,'7. 시연 경로')
bullet(d,'랜딩 페이지: http://localhost:4182/mulsallim-public-v2.html')
bullet(d,'실사용 대시보드: http://localhost:4182/mulsallim-dashboard-practice.html')
bullet(d,'권장 시연: 충청남도 → 논산시 → 탑정 선택 → 관측 근거 확인 → 이번 주 행동·현장 메모 → 브리핑 공유 → 재확인 일정 저장')
para(d,'공식 참고: 한국농어촌공사 농촌용수 저수지 수위정보(공공데이터포털), 한국농어촌공사 농업가뭄관리시스템 ADMS.',size=9,color='5D7470',before=8)

d.save(OUT)
print(OUT)
