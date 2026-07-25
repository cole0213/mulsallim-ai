from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

root = Path(__file__).parent
source = root / 'docs' / '05_최종_통합_서비스_명세_및_한글원고.md'
out_dir = root / 'deliverables'; out_dir.mkdir(exist_ok=True)
out = out_dir / '물살림AI_서비스개발_참가신청서_제출원고.docx'
doc = Document(); sec = doc.sections[0]
sec.top_margin = sec.bottom_margin = Cm(2.0); sec.left_margin = sec.right_margin = Cm(2.1)
style = doc.styles['Normal']; style.font.name = '맑은 고딕'; style._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕'); style.font.size = Pt(10.5)
title = doc.add_paragraph(); title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run('물살림 AI'); r.bold = True; r.font.size = Pt(24); r.font.color.rgb = RGBColor(20,63,65)
sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.add_run('가뭄 전환점 예측 및 공동 절수 시나리오 서비스\n서비스개발 분야 제출원고').font.size = Pt(12)
doc.add_paragraph('')
for raw in source.read_text(encoding='utf-8').splitlines():
    line = raw.strip()
    if not line or line.startswith('>'): continue
    if line.startswith('# '): continue
    if line.startswith('## '):
        p = doc.add_paragraph(); p.style = 'Heading 1'; p.add_run(line[3:]); continue
    if line.startswith('### '):
        p = doc.add_paragraph(); p.style = 'Heading 2'; p.add_run(line[4:]); continue
    if line.startswith('- '):
        p = doc.add_paragraph(style='List Bullet'); p.add_run(line[2:]); continue
    if line[:2].isdigit() and '. **' in line:
        p = doc.add_paragraph(); p.add_run(line).bold = False; continue
    if line.startswith('|') or line.startswith('```'): continue
    doc.add_paragraph(line.replace('**','').replace('`',''))
footer = sec.footer.paragraphs[0]; footer.alignment = WD_ALIGN_PARAGRAPH.CENTER; footer.text = '물살림 AI · 한국농어촌공사 공공데이터 활용 서비스개발 제안서'
doc.core_properties.title = '물살림 AI 서비스개발 참가신청서 제출원고'
doc.save(out)
print(out)
