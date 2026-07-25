from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).parent
OUT = ROOT / 'deliverables' / '물살림AI_대상권_최종제출원고_v7.docx'

def set_font(run, size=None, bold=None, color=None):
    run.font.name = '맑은 고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    if size: run.font.size = Pt(size)
    if bold is not None: run.bold = bold
    if color: run.font.color.rgb = RGBColor(*color)

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    el = OxmlElement('w:shd'); el.set(qn('w:fill'), fill); tcPr.append(el)

def cell_text(cell, text, bold=False, color=None):
    cell.text = ''
    p = cell.paragraphs[0]
    r = p.add_run(text); set_font(r, 9.5, bold, color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

def heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text); set_font(r, 15 if level == 1 else 12, True, (16,70,65))
    return p

def bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    set_font(p.add_run(text), 10.2)

def para(doc, text, bold_prefix=None):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(6)
    if bold_prefix and text.startswith(bold_prefix):
        set_font(p.add_run(bold_prefix), 10.2, True, (16,70,65))
        set_font(p.add_run(text[len(bold_prefix):]), 10.2)
    else: set_font(p.add_run(text), 10.2)
    return p

def table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers)); t.alignment = WD_TABLE_ALIGNMENT.CENTER; t.style='Table Grid'
    for i, h in enumerate(headers):
        cell_text(t.rows[0].cells[i], h, True, (255,255,255)); shade(t.rows[0].cells[i], '104641')
    for row in rows:
        cells=t.add_row().cells
        for i, v in enumerate(row):
            cell_text(cells[i], v)
            if i == 0: shade(cells[i], 'EAF4EF')
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths): row.cells[i].width=Cm(w)
    doc.add_paragraph().paragraph_format.space_after=Pt(0)
    return t

doc=Document()
sec=doc.sections[0]
sec.top_margin=Cm(1.8); sec.bottom_margin=Cm(1.8); sec.left_margin=Cm(1.9); sec.right_margin=Cm(1.9)
styles=doc.styles
styles['Normal'].font.name='맑은 고딕'; styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'),'맑은 고딕'); styles['Normal'].font.size=Pt(10.2)

p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(28)
r=p.add_run('제3회 KRC AI 디지털 혁신 공모전\n서비스개발 분야 참가신청서 첨부 원고'); set_font(r, 13, True, (8,122,119))
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('물살림 AI'); set_font(r, 30, True, (16,70,65))
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('코드 없는 지역 선택에서 근거 기반 현장 협의까지'); set_font(r, 13, True, (16,70,65))
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
set_font(p.add_run('KRC 공공데이터 기반 농업용수 의사결정 보조 서비스'), 10.5, False, (88,113,108))

heading(doc, '1. 서비스 개요')
para(doc, '물살림 AI는 농업인과 현장 담당자가 저수지 코드를 알 필요 없이 시·도와 시·군·구를 선택해 우리 지역 농업용수 변화를 확인하고, 위험 근거를 바탕으로 현장 확인·협의까지 연결하는 웹 서비스다.')
para(doc, '단순 수위 조회 화면에서 멈추지 않는다. 최신 저수율, 7일 변화, 관측 연속성을 함께 제시하고, 사용자가 다음 확인 행동을 선택한 뒤 개인정보 없이 30일짜리 현장 협의 카드를 공유하도록 설계했다.')
table(doc, ['구분', '내용'], [
    ['서비스명', '물살림 AI'],
    ['공모 분야', '서비스개발 — KRC 공공데이터와 AI를 활용한 Web-App 서비스'],
    ['핵심 사용자', '농업인·수리계, KRC 지역 현장 담당자, 지자체 농업용수 담당자'],
    ['핵심 가치', '코드 없는 지역 진입 / 근거가 보이는 변화 확인 / 사람 중심의 현장 협의'],
    ['공공데이터', '한국농어촌공사 농촌용수 저수지 수위정보 REST API'],
], [3.2, 12.8])

heading(doc, '2. 추진 배경과 문제 정의')
bullet(doc, '농업용수 정보는 저수지 코드, 관측일, 저수율처럼 현장 사용자가 바로 해석하기 어려운 형식으로 제공되는 경우가 많다.')
bullet(doc, '저수율의 숫자 하나만으로는 급격한 변화인지, 관측이 충분한지, 무엇을 먼저 확인해야 하는지 판단하기 어렵다.')
bullet(doc, '가뭄·급수 전환 판단은 공식 절차와 현장 확인이 필수지만, 그 전에 같은 근거를 빠르게 공유할 가벼운 협의 도구가 부족하다.')
para(doc, '따라서 본 서비스는 “예측을 과장하는 자동 판단”이 아니라, 공공데이터를 쉽게 찾게 하고 변화의 근거를 설명하며 현장 협의가 남도록 돕는 의사결정 보조에 집중한다.')

heading(doc, '3. 사용자 흐름과 주요 기능')
table(doc, ['단계', '사용자 경험', '구현 근거'], [
    ['① 지역 선택', '시·도 → 실제 관측이 있는 시·군·구 → 저수지 순으로 선택한다.', 'KRC county 조회 결과를 이용해 최근 관측 대상 지역만 동적으로 제시한다.'],
    ['② 변화 확인', '최신 저수율·7일 변화·관측일·추세 그래프·위험 근거를 한 화면에서 본다.', '그래프 툴팁과 축, 색상 위험표시, 키보드 탐색을 제공한다.'],
    ['③ 다음 행동', '관측 재확인, 급수 여력 확인, 다음 갱신일 확인 등 현재 단계에 맞는 행동을 본다.', '공식 지시가 아닌 현장 확인 우선순위로 표현한다.'],
    ['④ 협의 카드', '개인정보 없이 같은 근거와 현장 확인 상태를 공유한다.', '공개 링크는 읽기 전용이며 상태 변경은 생성자 비공개 키로만 가능하다.'],
    ['⑤ 재방문', '최근 조회·즐겨찾기·다음 확인 일정을 이용해 다시 들어온다.', '브라우저 저장소와 일정 파일(.ics)로 개인 작업 흐름을 지원한다.'],
], [2.1, 6.4, 7.5])

heading(doc, '4. AI 활용성과 책임 있는 배포 원칙')
para(doc, '현재 화면의 전환 신호는 최신 저수율, 단기 변화, 관측 연속성을 결합한 설명 가능한 분석 신호다. 사용자는 어떤 근거로 주의 단계가 제시됐는지 화면에서 확인할 수 있다.', '현재 화면의 전환 신호는')
table(doc, ['항목', '내용'], [
    ['오프라인 탐색', '탑정저수지 365일 관측 이력으로 저수율·7일 기울기·일별 변동성을 이용한 Isolation Forest 이상 전환 패턴 탐색을 수행했다.'],
    ['의미', '최근 패턴이 과거 관측 패턴 중 얼마나 드문지 보는 보조 근거이며, 미래 가뭄이나 저수율을 예측하는 모델이 아니다.'],
    ['배포 보류', '7일 후 저수율 예측 후보는 기준선보다 성능이 낮아 서비스 기능으로 배포하지 않았다.'],
    ['안전장치', '공식 경보 대체 금지, 자동 운영 결정 금지, 원자료·관측일 표시, 사람의 현장 확인과 상태 갱신을 적용한다.'],
], [3.0, 13.0])
para(doc, '이 원칙은 기술을 약하게 보이게 하기 위한 것이 아니라, 검증되지 않은 AI 결과가 농업용수 판단에 혼선을 주지 않도록 하는 서비스 품질 기준이다.')

heading(doc, '5. 공공데이터 활용·보안·개인정보 설계')
bullet(doc, 'KRC API 인증키는 브라우저나 제출 문서에 포함하지 않고 서버 환경변수에서만 사용한다.')
bullet(doc, '위치 조회는 최대 31일, 개별 저수지 상세는 최대 366일로 입력 범위를 제한한다.')
bullet(doc, 'API 요청 제한(분당 120회/IP), 5분 캐시, 12초 타임아웃을 적용해 외부 API 의존성을 관리한다.')
bullet(doc, '협의 카드는 이름·전화번호·정확한 농지 주소를 입력하지 않도록 안내하고, 30일 후 자동 삭제한다.')
bullet(doc, '공개 카드 URL에는 상태 변경 권한을 넣지 않는다. 생성자만 가진 비공개 키가 있어야 상태를 갱신할 수 있다.')

heading(doc, '6. 사업화 및 파일럿 계획')
table(doc, ['구분', '실행 계획'], [
    ['파일럿 대상', 'KRC 지역부서 1곳, 수리계/농업인 대표 10~20명, 4주'],
    ['검증 시나리오', '변화 신호 확인 → 관측 재확인 → 급수 여력·현장 상태 협의 카드 기록'],
    ['핵심 KPI', '첫 결과 도달 시간, 카드 생성·확인 완료율, 확인까지 걸린 시간, 불필요 경보율, 사용자 이해도'],
    ['성공 기준', '첫 결과 도달 중앙값 2분 이내, 카드 확인 완료율 70% 이상, 근거·주의문 이해도 80% 이상'],
    ['확장 방향', '강수·수요·급수·가뭄 이력 결합, 시간 분리 검증 후 제한적 예측 기능 검토, 지역 운영 대시보드 연계'],
], [3.0, 13.0])
para(doc, '본 원고는 파일럿 계획을 제시하며, 파일럿을 이미 완료한 것처럼 주장하지 않는다.')

heading(doc, '7. 기대효과')
bullet(doc, '국민·농업인이 저수지 코드나 API 문서를 몰라도 지역 선택만으로 농업용수 상황에 도달한다.')
bullet(doc, '숫자 조회를 변화 근거와 다음 행동으로 번역해 현장 확인의 시작 시간을 단축한다.')
bullet(doc, '개인정보를 쌓지 않는 단기 협의 카드로 농업인·현장 담당자 간 같은 근거의 소통을 지원한다.')
bullet(doc, '성능이 부족한 예측 모델을 배포하지 않는 원칙으로 공공 AI의 신뢰와 책임성을 확보한다.')

heading(doc, '8. 시연 안내')
para(doc, '지역 선택 화면: http://localhost:4184/  |  실사용 대시보드: mulsallim-dashboard-county-v7.html  |  AI 모델카드: mulsallim-ai-model-card.html')
para(doc, '시연에서는 제주특별자치도 등 실제 관측 지역을 선택해 저수지 코드 비노출, 변화 근거, 협의 카드의 공개/관리 권한 분리, AI 모델카드의 배포 보류 원칙을 순서대로 보여준다.')

heading(doc, '9. 최종 자체 검증')
para(doc, '2026-07-24 기준 로컬 v7 회귀 테스트에서 KRC 건강 상태, 제주 지역 조회 및 개별 이력, 서울 무관측 처리, 주요 화면 HTTP 응답, 협의 카드 관리 키 비노출, 무권한 상태 변경 차단을 확인했다.')

OUT.parent.mkdir(exist_ok=True)
doc.save(OUT)
print(OUT)
