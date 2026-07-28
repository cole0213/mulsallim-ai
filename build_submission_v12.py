from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

ROOT = Path(__file__).parent
SRC = ROOT / 'deliverables' / '물살림AI_대상권_실증·안전성_최종제출원고_v11.docx'
OUT = ROOT / 'deliverables' / '물살림AI_대상권_공개배포완료_최종제출원고_v12.docx'


def font(run, size=10.2, bold=None):
    run.font.name = '맑은 고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(6)
    font(p.add_run(text), 15, True)


def para(doc, text, size=10.2):
    p = doc.add_paragraph()
    font(p.add_run(text), size)


def bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    font(p.add_run(text), 10.2)


d = Document(SRC)

heading(d, '15. 공개 HTTPS 배포 완료 (2026-07-27)')
para(d, '물살림 AI는 고정 HTTPS 공개 주소에서 지역 선택부터 저수율·AI 전환 신호·현장 협의까지 전 흐름이 실제로 동작한다. 상위망이 80·443 인입을 차단해 오리진 직접 인증서 발급이 불가능하던 문제를 Cloudflare Pages 정적 호스팅과 서버 간 프록시로 해결했다.')
bullet(d, '공개 주소: https://mulsallim-ai.pages.dev  (진입 mulsallim-public-county-v12.html)')
bullet(d, '대시보드: mulsallim-dashboard-county-v12.html  |  다지역 모델카드: mulsallim-ai-model-card-v3.html  |  파일럿 피드백: mulsallim-pilot.html')
bullet(d, '구조: 브라우저는 Pages HTTPS만 사용하고, /api/* 요청은 Pages Function이 KRC 운영 서버로 서버 간 프록시한다. KRC 인증키는 운영 서버 환경변수에만 두고 Pages·GitHub·클라이언트에 노출하지 않는다.')
bullet(d, '보안: 정적 자산에 CSP(스크립트 eval 차단)·X-Frame-Options·Referrer-Policy·Permissions-Policy를 적용하고, 배포 산출물에서 실행 코드의 eval 사용을 제거했다.')
bullet(d, '검증: 9개 시·도 15개 이상 저수지에서 지역 조회·상세 관측·AI 신호·지역 비교가 오류 없이 렌더됨을 확인했다.')

heading(d, '16. 이번 개선으로 강화된 기능')
bullet(d, 'AI 예측 궤적 시각화: 저수율 그래프에 최근 추세를 이어 관심선(50%)·주의선(35%) 도달 예상 시점을 점선 궤적과 교차점(예: 관심 D+10, 주의 D+18)으로 표시한다. 텍스트 도달일과 수치가 일치하며 확정 예측이 아님을 명시한다.')
bullet(d, '지역 확인 우선순위: 같은 시·군의 저수지를 AI 전환 신호 순으로 정렬해 "지금 가장 먼저 확인할 곳"을 문장과 표로 제시한다. 어느 저수지·수혜구역부터 점검·협의할지에 답한다.')
bullet(d, '오류 복구: 데이터 조회 실패 시 재시도 버튼을 제공하고, 지역 목록·저수지 조회 실패를 사용자 언어로 안내한다.')
bullet(d, '접근성: 건너뛰기 링크, 포커스 표시, ARIA 상태 안내, 그래프의 키보드 화살표 낭독(날짜별 수치)을 제공한다.')
bullet(d, '완성도: 모바일 360px 무결, 카드 깊이·차트 그라디언트·호버 크로스헤어로 데이터 화면의 가독성과 완성도를 높였다.')

heading(d, '17. 정직성 원칙 (변경 없음)')
para(d, '성능 검증에서 기준을 넘지 못한 예측 모델은 배포하지 않는다. 실제 사용자 파일럿 응답과 협의 카드 완료는 아직 확보 전이며 성과 지표로 주장하지 않는다. 화면·문서·발표는 같은 기능과 같은 공개 주소를 가리킨다.')
para(d, '운영 안내: 공개 데모의 실시간 데이터는 KRC 운영 서버와의 프록시 연결이 활성일 때 제공된다. 심사·시연 시 운영 서버 연결을 사전 점검한다.')

d.save(OUT)
print(OUT)
