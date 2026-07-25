from pathlib import Path
from docx import Document

root=Path(__file__).parent
source=root/'deliverables'/'물살림AI_실사용중심_최종제출원고_최종v3.docx'
output=root/'deliverables'/'물살림AI_실사용중심_최종제출원고_운영형v4.docx'
d=Document(source)
for p in d.paragraphs:
    for r in p.runs:
        r.text=r.text.replace('http://localhost:4183/mulsallim-public-county-v3.html','http://localhost:4184/mulsallim-public-county-v4.html').replace('http://localhost:4183/mulsallim-dashboard-county-v3.html','http://localhost:4184/mulsallim-dashboard-county-v4.html')
for t in d.tables:
    for row in t.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:r.text=r.text.replace('http://localhost:4183/mulsallim-public-county-v3.html','http://localhost:4184/mulsallim-public-county-v4.html').replace('http://localhost:4183/mulsallim-dashboard-county-v3.html','http://localhost:4184/mulsallim-dashboard-county-v4.html')
d.save(output)
print(output)
